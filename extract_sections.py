"""Extract thesis sections — find the FIRST occurrence of each heading AFTER the TOC."""
import os, re
from docx import Document

DOCX_PATH = "/Users/fc/Nutstore Files/我的坚果云/毕业/毕业论文/毕业论文+作者简介.docx"
OUTPUT_DIR = "/Users/fc/Documents/thesis-translation"

doc = Document(DOCX_PATH)

# Find the TOC range: from the first "toc" style paragraph after "目录" through the last consecutive toc entry
toc_range = range(0, 0)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().replace(" ", "") == "目录":
        # Find the first toc-style paragraph after this
        first_toc = None
        for j in range(i+1, len(doc.paragraphs)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if "toc" in style.lower():
                first_toc = j
                break
        if first_toc is None:
            break
        # Now find the last consecutive toc-style paragraph
        last_toc = first_toc
        for j in range(first_toc+1, len(doc.paragraphs)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if "toc" in style.lower():
                last_toc = j
            else:
                break
        toc_range = range(first_toc, last_toc + 1)
        print(f"TOC range: paragraphs {first_toc} to {last_toc}")
        break

# After the TOC, find the first occurrence of each heading marker
section_markers = [
    ("致 谢",       "00_acknowledgements"),
    ("摘要",        "01_abstract_cn"),
    ("Abstract",    "02_abstract_en"),
    ("绪论",        "03_introduction"),
    ("实验材料与方法", "04_methods"),
    ("实验结果",    "05_results"),      # can't match loosely or "实验结果" in discussion body triggers false+
    ("实验结论",    "06_conclusion"),
    ("讨论",        "07_discussion"),
    ("参考文献",    "08_references"),
    ("作者简历",    "09_author_bio"),
]

# For markers that might appear in prose, use strict matching:
# The actual section heading in the body is typically just the number + title,
# e.g. "3 实验结果", "4 实验结论", "5 讨论"
# But in prose these words may appear too.
# Strategy: require the marker to appear at the START of the paragraph text
# and the paragraph should have meaningful content before any line break.

def is_section_heading(text, marker):
    """Check if para text looks like a real section heading."""
    stripped = text.strip()
    if marker == "Abstract":
        return stripped == "Abstract"
    if marker == "致 谢" or marker == "摘要":
        return stripped == marker or stripped == marker.replace(" ", "")
    if marker == "实验材料与方法":
        # This can be "2 实验材料与方法" or just "实验材料与方法"
        return "实验材料与方法" in stripped
    if marker == "实验结果":
        return re.match(r'^\s*\d+\s+实验结果', stripped) is not None
    if marker == "实验结论":
        return re.match(r'^\s*\d+\s+实验结论', stripped) is not None
    if marker == "讨论":
        return re.match(r'^\s*\d+\s+讨论', stripped) is not None
    if marker == "参考文献":
        return re.match(r'^\s*\d+\s+参考文献', stripped) is not None
    if marker == "绪论":
        return re.match(r'^\s*\d*\s*绪论', stripped) is not None
    if marker == "作者简历":
        return stripped == "作者简历"
    return marker in stripped

starts = {}
for marker_text, section_key in section_markers:
    for i, para in enumerate(doc.paragraphs):
        if i in toc_range:
            continue
        text = para.text.strip()
        if is_section_heading(text, marker_text):
            starts[section_key] = i
            print(f"  {section_key} @ para {i}: '{text[:60]}'")
            break

# Sort by paragraph index
sorted_keys = sorted(starts, key=lambda k: starts[k])

# Write sections
file_names = {
    "00_acknowledgements": "00_acknowledgements.md",
    "01_abstract_cn": "01_abstract_cn.md",
    "02_abstract_en": "02_abstract_en.md",
    "03_introduction": "03_introduction.md",
    "04_methods": "04_methods.md",
    "05_results": "05_results.md",
    "06_conclusion": "06_conclusion.md",
    "07_discussion": "07_discussion.md",
    "08_references": "08_references.md",
    "09_author_bio": "09_author_bio.md",
}

sections_dir = os.path.join(OUTPUT_DIR, "sections")
os.makedirs(sections_dir, exist_ok=True)

section_stats = []
for idx, key in enumerate(sorted_keys):
    start = starts[key]
    end = starts[sorted_keys[idx + 1]] if idx + 1 < len(sorted_keys) else len(doc.paragraphs)

    paras = []
    char_count = 0
    for i in range(start, end):
        t = doc.paragraphs[i].text
        if t.strip():
            para_text = t.strip()
            paras.append(para_text)
            char_count += len(para_text)

    content = "\n\n".join(paras)
    filepath = os.path.join(sections_dir, file_names[key])
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    section_stats.append({
        "section": key,
        "file": file_names[key],
        "chars": char_count,
        "lines": len(paras),
    })
    print(f"  Written: {file_names[key]} ({char_count:,} chars, {len(paras)} lines)")

# Inventory
total_chars = sum(s["chars"] for s in section_stats)
inventory = f"""# Thesis Section Inventory

Source: `毕业论文+作者简介.docx`
Extracted: 2026-05-30

| # | Section | File | Characters | Lines (non-blank) |
|---|---------|------|----------:|------------------:|
"""
for s in section_stats:
    inventory += f"| {s['section']} | {s['file']} | {s['chars']:,} | {s['lines']} |\n"
inventory += f"\n**Total**: {total_chars:,} chars across {len(section_stats)} sections\n"

with open(os.path.join(OUTPUT_DIR, "section_inventory.md"), "w", encoding="utf-8") as f:
    f.write(inventory)
print(f"\n=== Section inventory written ===")
print(inventory)